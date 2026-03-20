"""
Document Loader — Ingestion Layer
===================================
Supports PDF, plain text, Markdown, and DOCX.
Returns a uniform list of RawDocument objects regardless of source format.

Design decisions:
- Pydantic models enforce schema at the boundary
- PyMuPDF (fitz) for PDF: fast, accurate, handles scanned PDFs with optional OCR
- python-docx for DOCX: handles tables, headers, footers
- Each loader is a separate class implementing DocumentLoaderBase for easy extension
"""

from __future__ import annotations

import hashlib
import os
from abc import ABC, abstractmethod
from datetime import datetime, timezone
from pathlib import Path
from typing import List, Optional

from pydantic import BaseModel, Field

from config.settings import settings
from observability.logger import get_logger

logger = get_logger(__name__)


# ─────────────────────────────────────────────────────────────
# Domain models
# ─────────────────────────────────────────────────────────────

class DocumentMetadata(BaseModel):
    """Metadata associated with a loaded document."""
    source: str                          # File path or URL
    filename: str
    doc_type: str                        # pdf, txt, md, docx
    file_size_bytes: int
    content_hash: str                    # SHA-256 for deduplication
    loaded_at: str = Field(
        default_factory=lambda: datetime.now(timezone.utc).isoformat()
    )
    page_count: Optional[int] = None
    author: Optional[str] = None
    title: Optional[str] = None
    extra: dict = Field(default_factory=dict)


class RawDocument(BaseModel):
    """A fully loaded document ready for chunking."""
    doc_id: str                          # Derived from content hash
    content: str                         # Full extracted text
    metadata: DocumentMetadata

    class Config:
        arbitrary_types_allowed = True


def _compute_hash(content: str) -> str:
    return hashlib.sha256(content.encode("utf-8")).hexdigest()[:16]


# ─────────────────────────────────────────────────────────────
# Loader interface
# ─────────────────────────────────────────────────────────────

class DocumentLoaderBase(ABC):
    """Abstract base — all loaders implement load()."""

    @abstractmethod
    def load(self, file_path: Path) -> RawDocument:
        raise NotImplementedError

    def _build_doc(
        self,
        file_path: Path,
        content: str,
        doc_type: str,
        extra_metadata: Optional[dict] = None,
    ) -> RawDocument:
        content_hash = _compute_hash(content)
        stat = file_path.stat()
        metadata = DocumentMetadata(
            source=str(file_path.resolve()),
            filename=file_path.name,
            doc_type=doc_type,
            file_size_bytes=stat.st_size,
            content_hash=content_hash,
            extra=extra_metadata or {},
        )
        return RawDocument(
            doc_id=f"{doc_type[:3]}-{content_hash}",
            content=content,
            metadata=metadata,
        )


# ─────────────────────────────────────────────────────────────
# PDF Loader
# ─────────────────────────────────────────────────────────────

class PDFLoader(DocumentLoaderBase):
    """
    Extracts text from PDFs using PyMuPDF (fitz).
    Falls back to pytesseract OCR for image-based pages.
    """

    def load(self, file_path: Path) -> RawDocument:
        try:
            import fitz  # PyMuPDF
        except ImportError:
            raise ImportError("PyMuPDF not installed: pip install pymupdf")

        doc = fitz.open(str(file_path))
        pages_text: List[str] = []
        page_count = len(doc)

        for page_num, page in enumerate(doc):
            text = page.get_text("text")
            if not text.strip():
                # Image-based page — attempt OCR
                text = self._ocr_page(page, page_num)
            pages_text.append(text)

        doc.close()
        full_text = "\n\n".join(pages_text)

        logger.info(
            f"PDF loaded: {file_path.name}",
            extra={"pages": page_count, "chars": len(full_text)},
        )

        raw_doc = self._build_doc(
            file_path, full_text, "pdf", {"page_count": page_count}
        )
        raw_doc.metadata.page_count = page_count
        return raw_doc

    @staticmethod
    def _ocr_page(page: "fitz.Page", page_num: int) -> str:
        """OCR fallback using pytesseract. Returns empty string if unavailable."""
        try:
            import pytesseract
            from PIL import Image
            import io

            pix = page.get_pixmap(dpi=200)
            img = Image.open(io.BytesIO(pix.tobytes("png")))
            return pytesseract.image_to_string(img)
        except ImportError:
            logger.debug(f"OCR unavailable for page {page_num}, skipping")
            return ""


# ─────────────────────────────────────────────────────────────
# Plain Text / Markdown Loader
# ─────────────────────────────────────────────────────────────

class TextLoader(DocumentLoaderBase):
    """Handles .txt and .md files."""

    def load(self, file_path: Path) -> RawDocument:
        encoding = "utf-8"
        try:
            content = file_path.read_text(encoding=encoding)
        except UnicodeDecodeError:
            content = file_path.read_text(encoding="latin-1")

        doc_type = file_path.suffix.lstrip(".") or "txt"
        logger.info(f"Text loaded: {file_path.name} ({len(content)} chars)")
        return self._build_doc(file_path, content, doc_type)


# ─────────────────────────────────────────────────────────────
# DOCX Loader
# ─────────────────────────────────────────────────────────────

class DOCXLoader(DocumentLoaderBase):
    """
    Extracts text from Word documents including tables.
    Tables are rendered as tab-separated rows for context preservation.
    """

    def load(self, file_path: Path) -> RawDocument:
        try:
            from docx import Document
        except ImportError:
            raise ImportError("python-docx not installed: pip install python-docx")

        doc = Document(str(file_path))
        parts: List[str] = []

        # Paragraphs
        for para in doc.paragraphs:
            if para.text.strip():
                parts.append(para.text)

        # Tables — flatten into readable text
        for table in doc.tables:
            for row in table.rows:
                row_text = "\t".join(cell.text.strip() for cell in row.cells)
                if row_text.strip():
                    parts.append(row_text)

        content = "\n\n".join(parts)
        logger.info(f"DOCX loaded: {file_path.name} ({len(parts)} text blocks)")
        return self._build_doc(file_path, content, "docx")


# ─────────────────────────────────────────────────────────────
# Document Loader Factory
# ─────────────────────────────────────────────────────────────

_LOADER_REGISTRY: dict = {
    ".pdf": PDFLoader,
    ".txt": TextLoader,
    ".md": TextLoader,
    ".docx": DOCXLoader,
}


class DocumentLoader:
    """
    Main entry point for document loading.
    Auto-detects format and dispatches to the correct loader.
    """

    def __init__(self) -> None:
        self._loaders = {ext: cls() for ext, cls in _LOADER_REGISTRY.items()}

    def load(self, file_path: str | Path) -> RawDocument:
        path = Path(file_path)

        if not path.exists():
            raise FileNotFoundError(f"Document not found: {path}")

        size_mb = path.stat().st_size / (1024 * 1024)
        if size_mb > settings.MAX_DOCUMENT_SIZE_MB:
            raise ValueError(
                f"Document too large: {size_mb:.1f} MB > {settings.MAX_DOCUMENT_SIZE_MB} MB limit"
            )

        ext = path.suffix.lower()
        loader = self._loaders.get(ext)
        if not loader:
            raise ValueError(
                f"Unsupported format: {ext}. "
                f"Supported: {list(self._loaders.keys())}"
            )

        return loader.load(path)

    def load_directory(
        self, directory: str | Path, recursive: bool = True
    ) -> List[RawDocument]:
        """
        Load all supported documents from a directory.
        Skips unsupported extensions silently.
        """
        dir_path = Path(directory)
        if not dir_path.is_dir():
            raise NotADirectoryError(f"Not a directory: {dir_path}")

        pattern = "**/*" if recursive else "*"
        documents: List[RawDocument] = []

        for file_path in dir_path.glob(pattern):
            if file_path.suffix.lower() in self._loaders:
                try:
                    doc = self.load(file_path)
                    documents.append(doc)
                except Exception as exc:
                    logger.warning(
                        f"Failed to load {file_path.name}: {exc}"
                    )

        logger.info(
            f"Directory loaded: {dir_path.name}",
            extra={"docs_loaded": len(documents)},
        )
        return documents
